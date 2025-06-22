from docx import Document
from docx.shared import Pt
from fpdf import FPDF
import os
import re

def save_to_docx(text: str, name: str) -> str:
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    for para_text in text.split('\n\n'):
        p = doc.add_paragraph()
        parts = re.split(r'(\*\*.*?\*\*)', para_text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                p.add_run(part[2:-2]).bold = True
            else:
                p.add_run(part)
    path = f"{name}.docx"
    doc.save(path)
    return path

def save_to_pdf(text: str, name: str) -> str:
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    font_path_regular = "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"
    font_path_bold = "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"
    pdf.add_font("DejaVu", "", font_path_regular, uni=True)
    pdf.add_font("DejaVu", "B", font_path_bold, uni=True)
    pdf.set_font("DejaVu", size=11)
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            pdf.set_font("DejaVu", "B", size=11)
            pdf.write(5, part[2:-2])
        else:
            pdf.set_font("DejaVu", "", size=11)
            pdf.write(5, part)
    path = f"{name}.pdf"
    pdf.output(path)
    return path

def save_to_py(text: str, name: str) -> str:
    path = f"{name}.py"
    with open(path, "w", encoding="utf-8") as f:
        f.write(text)
    return path

def split_text(text: str, chunk_size: int = 4096) -> list[str]:
    if len(text) <= chunk_size:
        return [text]
    chunks = []
    while text:
        if len(text) <= chunk_size:
            chunks.append(text)
            break
        split_pos = text.rfind('\n', 0, chunk_size)
        if split_pos == -1:
            split_pos = chunk_size        
        chunks.append(text[:split_pos])
        text = text[split_pos:].lstrip()
    return chunks

def save_code_to_file(text: str, name: str, language: str) -> str:
    """Сохраняет текстовое содержимое в файл с нужным расширением."""
    extension = "py" if language == "python" else "java"
    path = f"{name}.{extension}"
    with open(path, "w", encoding="utf-8") as f:
        f.write(text)
    return path